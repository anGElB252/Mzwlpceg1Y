# 代码生成时间: 2025-08-25 01:48:16
import dash
import dash_core_components as dcc
import dash_html_components as html
import dash_bootstrap_components as dbc
from dash.dependencies import Input, Output, State
from docx import Document
from docx.oxml.ns import qn
import io
from flask import send_file
import os

# Define the app
app = dash.Dash(__name__, external_stylesheets=[dbc.themes.BOOTSTRAP])

# Define the layout
app.layout = dbc.Container([
    dbc.Row([
        dbc.Col(html.H1("Document Converter"), width=12)
    ]),
    dbc.Row([
        dbc.Col(dcc.Upload(
            id='upload-data',
            children=html.Div([
                'Drag and Drop or ',
                html.A('Select a File')
            ]),
            style={'width': '100%', 'height': '60px', 'lineHeight': '60px',
                   'borderWidth': '1px', 'borderStyle': 'dashed',
                   'borderRadius': '5px', 'textAlign': 'center',
                   'margin': '10px'},
            multiple=True
        ), width=12)
    ]),
    dbc.Row([
        dbc.Col(html.Div(id='output-data-upload-container'), width=12)
    ])
])

# Callback to convert and display the document
@app.callback(
    Output('output-data-upload-container', 'children'),
    [Input('upload-data', 'contents')],
    [State('upload-data', 'filename'), State('upload-data', 'last_modified')]
)
def update_output(contents, filename, last_modified):
    if contents is not None:
        # Create an in-memory bytes buffer
        output = io.BytesIO()
        doc = Document()
        for content in contents:
            filename = content['filename']
            last_modified = content['last_modified']
            # Read the file content and add to the document
            doc.add_paragraph(filename)
            doc.add_paragraph(str(last_modified))
            # Save the document to the buffer
            doc.save(output)
            output.seek(0)
        return html.Div([
            html.Hr(),
            html.P(f'Filename: {filename}'),
            html.P(f'Last Modified: {last_modified}'),
            html.A(dbc.Button("Download", color="primary", className="mr-2"),
                   href=f'/download?filename={filename}'),
            html.A(dbc.Button("Download", color="primary"),
                   href=f'/download?filename={filename}')
        ])
    return html.Div()

# Route for downloading the converted document
@app.server.route("/download")
def download():
    filename = dash.server.request.args.get("filename")
    if filename:
        file_path = os.path.join("./downloads", filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
    return "File not found", 404

# Run the server
if __name__ == '__main__':
    app.run_server(debug=True)
